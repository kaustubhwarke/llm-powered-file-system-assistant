"""Core file system tools exposed to the LLM (Part A).

Each public function:
  * validates input via a Pydantic model,
  * resolves the path inside the configured sandbox root,
  * performs the operation,
  * returns a JSON-serializable ``dict`` derived from a Pydantic result model,
  * never raises on expected I/O failures — errors are reported in the result.

The tools are intentionally provider-neutral: the same dict-returning
contract is what the LLM tool-use loop consumes regardless of vendor.
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from pydantic import ValidationError

from llm_file_assistant.config import Settings, get_settings
from llm_file_assistant.exceptions import (
    FileNotFoundInSandboxError,
    FileParseError,
    FileSystemError,
    FileTooLargeError,
    PathSecurityError,
    UnsupportedFileTypeError,
)
from llm_file_assistant.logging_config import get_logger
from llm_file_assistant.schemas import (
    FileMetadata,
    ListFilesInput,
    ListFilesResult,
    ReadFileInput,
    ReadFileResult,
    SearchInFileInput,
    SearchInFileResult,
    SearchMatch,
    ToolStatus,
    WriteFileInput,
    WriteFileResult,
)

logger = get_logger(__name__)

SUPPORTED_READ_EXTENSIONS: frozenset[str] = frozenset({".txt", ".md", ".pdf", ".docx"})


# =============================================================================
# Internal helpers
# =============================================================================


def _resolve_inside_sandbox(relative: str, settings: Settings) -> Path:
    """Resolve a user-supplied relative path inside the sandbox root.

    Raises:
        PathSecurityError: if the resolved path escapes the sandbox root.
    """
    root = settings.fs_root
    root.mkdir(parents=True, exist_ok=True)
    candidate = (root / relative).expanduser()
    try:
        resolved = candidate.resolve()
    except OSError as exc:
        raise PathSecurityError(f"Cannot resolve path: {relative}") from exc

    try:
        resolved.relative_to(root)
    except ValueError as exc:
        raise PathSecurityError(
            f"Path '{relative}' escapes sandbox root '{root}'"
        ) from exc
    return resolved


def _build_metadata(path: Path, settings: Settings) -> FileMetadata:
    """Build a FileMetadata record for an existing file."""
    stat = path.stat()
    return FileMetadata(
        name=path.name,
        path=str(path.relative_to(settings.fs_root)).replace("\\", "/"),
        extension=path.suffix.lower(),
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
    )


def _error_dict(exc: Exception) -> dict[str, Any]:
    """Translate an exception into a result-envelope error dict."""
    return {
        "status": ToolStatus.ERROR.value,
        "error": str(exc),
        "error_type": exc.__class__.__name__,
    }


def _normalize_extension(ext: str | None) -> str | None:
    if ext is None:
        return None
    ext = ext.strip().lower()
    if not ext:
        return None
    return ext if ext.startswith(".") else f".{ext}"


# =============================================================================
# Document parsers
# =============================================================================


def _read_txt(path: Path) -> tuple[str, dict[str, Any]]:
    """Read a plain-text file with encoding fallbacks."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return path.read_text(encoding=encoding), {}
        except UnicodeDecodeError:
            continue
    raise FileParseError(f"Could not decode text file with known encodings: {path.name}")


def _read_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    """Read a PDF file and extract its text content."""
    try:
        from pypdf import PdfReader  # local import keeps cold start fast
        from pypdf.errors import PdfReadError
    except ImportError as exc:  # pragma: no cover - dependency guaranteed by requirements
        raise FileParseError("pypdf is not installed") from exc

    try:
        reader = PdfReader(BytesIO(path.read_bytes()))
        if reader.is_encrypted:
            raise FileParseError(f"PDF is encrypted: {path.name}")
        pages = [page.extract_text() or "" for page in reader.pages]
    except PdfReadError as exc:
        raise FileParseError(f"Malformed PDF '{path.name}': {exc}") from exc

    text = "\n\n".join(pages).strip()
    return text, {"page_count": len(pages)}


def _read_docx(path: Path) -> tuple[str, dict[str, Any]]:
    """Read a DOCX file and extract its text content (paragraphs + tables)."""
    try:
        from docx import Document  # local import keeps cold start fast
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:  # pragma: no cover
        raise FileParseError("python-docx is not installed") from exc

    try:
        document = Document(str(path))
    except PackageNotFoundError as exc:
        raise FileParseError(f"Malformed DOCX '{path.name}': {exc}") from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), {}


_PARSERS = {
    ".txt": _read_txt,
    ".md": _read_txt,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


# =============================================================================
# Public tools
# =============================================================================


def read_file(filepath: str) -> dict[str, Any]:
    """Read a file from the sandbox and return its text content.

    Supports plain text (``.txt``, ``.md``), PDF (``.pdf``), and Word
    (``.docx``) documents.

    Args:
        filepath: Path relative to the configured sandbox root.

    Returns:
        A JSON-serializable dict shaped like :class:`ReadFileResult`. The
        ``status`` field is ``"success"`` or ``"error"``. On success,
        ``content`` and ``metadata`` are populated; on error, ``error`` and
        ``error_type`` describe the failure.
    """
    settings = get_settings()
    try:
        params = ReadFileInput(filepath=filepath)
        path = _resolve_inside_sandbox(params.filepath, settings)

        if not path.exists():
            raise FileNotFoundInSandboxError(f"File not found: {params.filepath}")
        if not path.is_file():
            raise FileSystemError(f"Not a file: {params.filepath}")

        ext = path.suffix.lower()
        if ext not in _PARSERS:
            raise UnsupportedFileTypeError(
                f"Unsupported extension '{ext}'. Supported: "
                f"{sorted(SUPPORTED_READ_EXTENSIONS)}"
            )

        size = path.stat().st_size
        if size > settings.fs_max_file_bytes:
            raise FileTooLargeError(
                f"File '{params.filepath}' is {size} bytes "
                f"(limit {settings.fs_max_file_bytes})"
            )

        content, extra = _PARSERS[ext](path)
        result = ReadFileResult(
            status=ToolStatus.SUCCESS,
            content=content,
            metadata=_build_metadata(path, settings),
            page_count=extra.get("page_count"),
        )
        logger.info(
            "tool.read_file.ok",
            filepath=params.filepath,
            size_bytes=size,
            extension=ext,
        )
        return result.model_dump(mode="json")
    except ValidationError as exc:
        logger.warning("tool.read_file.invalid_input", error=str(exc))
        return _error_dict(exc)
    except FileSystemError as exc:
        logger.warning("tool.read_file.error", filepath=filepath, error=str(exc))
        return _error_dict(exc)
    except Exception as exc:  # noqa: BLE001 - boundary catch for LLM safety
        logger.exception("tool.read_file.unexpected", filepath=filepath)
        return _error_dict(exc)


def list_files(directory: str, extension: str | None = None) -> dict[str, Any]:
    """List files in a directory within the sandbox.

    Args:
        directory: Path relative to the configured sandbox root.
        extension: Optional file extension filter (e.g. ``"pdf"`` or ``".pdf"``).

    Returns:
        A JSON-serializable dict shaped like :class:`ListFilesResult`.
    """
    settings = get_settings()
    try:
        params = ListFilesInput(directory=directory, extension=extension)
        path = _resolve_inside_sandbox(params.directory, settings)

        if not path.exists():
            raise FileNotFoundInSandboxError(
                f"Directory not found: {params.directory}"
            )
        if not path.is_dir():
            raise FileSystemError(f"Not a directory: {params.directory}")

        ext_filter = _normalize_extension(params.extension)

        entries: list[FileMetadata] = []
        for entry in sorted(path.iterdir(), key=lambda p: p.name.lower()):
            if not entry.is_file():
                continue
            if ext_filter and entry.suffix.lower() != ext_filter:
                continue
            entries.append(_build_metadata(entry, settings))

        result = ListFilesResult(
            status=ToolStatus.SUCCESS,
            directory=str(path.relative_to(settings.fs_root)).replace("\\", "/") or ".",
            count=len(entries),
            files=entries,
        )
        logger.info(
            "tool.list_files.ok",
            directory=params.directory,
            extension=ext_filter,
            count=len(entries),
        )
        return result.model_dump(mode="json")
    except ValidationError as exc:
        logger.warning("tool.list_files.invalid_input", error=str(exc))
        return _error_dict(exc)
    except FileSystemError as exc:
        logger.warning("tool.list_files.error", directory=directory, error=str(exc))
        return _error_dict(exc)
    except Exception as exc:  # noqa: BLE001
        logger.exception("tool.list_files.unexpected", directory=directory)
        return _error_dict(exc)


def write_file(filepath: str, content: str) -> dict[str, Any]:
    """Write UTF-8 text content to a file within the sandbox.

    Parent directories are created automatically. If the target file already
    exists, it is overwritten (unless ``overwrite=False`` is passed via the
    underlying schema, which the LLM controls).

    Args:
        filepath: Path relative to the configured sandbox root.
        content: UTF-8 text to write.

    Returns:
        A JSON-serializable dict shaped like :class:`WriteFileResult`.
    """
    settings = get_settings()
    try:
        params = WriteFileInput(filepath=filepath, content=content)
        path = _resolve_inside_sandbox(params.filepath, settings)

        if path.exists() and not params.overwrite:
            raise FileSystemError(
                f"File exists and overwrite=false: {params.filepath}"
            )
        if path.exists() and path.is_dir():
            raise FileSystemError(f"Target is a directory: {params.filepath}")

        encoded = content.encode("utf-8")
        if len(encoded) > settings.fs_max_file_bytes:
            raise FileTooLargeError(
                f"Refusing to write {len(encoded)} bytes "
                f"(limit {settings.fs_max_file_bytes})"
            )

        created_dirs = not path.parent.exists()
        path.parent.mkdir(parents=True, exist_ok=True)
        overwritten = path.exists()
        path.write_bytes(encoded)

        result = WriteFileResult(
            status=ToolStatus.SUCCESS,
            path=str(path.relative_to(settings.fs_root)).replace("\\", "/"),
            bytes_written=len(encoded),
            created_directories=created_dirs,
            overwritten=overwritten,
        )
        logger.info(
            "tool.write_file.ok",
            filepath=params.filepath,
            bytes_written=len(encoded),
            overwritten=overwritten,
        )
        return result.model_dump(mode="json")
    except ValidationError as exc:
        logger.warning("tool.write_file.invalid_input", error=str(exc))
        return _error_dict(exc)
    except FileSystemError as exc:
        logger.warning("tool.write_file.error", filepath=filepath, error=str(exc))
        return _error_dict(exc)
    except Exception as exc:  # noqa: BLE001
        logger.exception("tool.write_file.unexpected", filepath=filepath)
        return _error_dict(exc)


def search_in_file(
    filepath: str, keyword: str, context_lines: int = 2
) -> dict[str, Any]:
    """Case-insensitive substring search inside a file.

    Args:
        filepath: Path relative to the configured sandbox root.
        keyword: Substring to search for (case-insensitive).
        context_lines: Number of surrounding lines to include with each match.

    Returns:
        A JSON-serializable dict shaped like :class:`SearchInFileResult`.
    """
    try:
        params = SearchInFileInput(
            filepath=filepath, keyword=keyword, context_lines=context_lines
        )
    except ValidationError as exc:
        logger.warning("tool.search_in_file.invalid_input", error=str(exc))
        return _error_dict(exc)

    read_result = read_file(params.filepath)
    if read_result.get("status") != ToolStatus.SUCCESS.value:
        logger.warning(
            "tool.search_in_file.read_failed",
            filepath=params.filepath,
            error=read_result.get("error"),
        )
        return read_result

    try:
        content = read_result.get("content") or ""
        needle = params.keyword.lower()
        lines = content.splitlines()
        matches: list[SearchMatch] = []
        for idx, line in enumerate(lines):
            offset = line.lower().find(needle)
            if offset == -1:
                continue
            start = max(0, idx - params.context_lines)
            end = min(len(lines), idx + params.context_lines + 1)
            matches.append(
                SearchMatch(
                    line_number=idx + 1,
                    line=line,
                    context_before=lines[start:idx],
                    context_after=lines[idx + 1 : end],
                    char_offset=offset,
                )
            )

        result = SearchInFileResult(
            status=ToolStatus.SUCCESS,
            filepath=params.filepath,
            keyword=params.keyword,
            match_count=len(matches),
            matches=matches,
        )
        logger.info(
            "tool.search_in_file.ok",
            filepath=params.filepath,
            keyword=params.keyword,
            match_count=len(matches),
        )
        return result.model_dump(mode="json")
    except Exception as exc:  # noqa: BLE001
        logger.exception("tool.search_in_file.unexpected", filepath=params.filepath)
        return _error_dict(exc)
