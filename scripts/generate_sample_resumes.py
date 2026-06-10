"""Generate 10 dummy resume files across .txt, .docx, and .pdf formats.

Idempotent — running it again overwrites the existing files. Used to
seed ``data/resumes/`` for demos and tests.
"""

from __future__ import annotations

import sys
from pathlib import Path

# Make src importable without an editable install.
PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))


RESUMES: list[dict[str, str]] = [
    {
        "filename": "alice_johnson.txt",
        "format": "txt",
        "name": "Alice Johnson",
        "title": "Senior Backend Engineer",
        "email": "alice.johnson@example.com",
        "phone": "+1-415-555-0142",
        "location": "San Francisco, CA",
        "summary": (
            "Senior backend engineer with 8+ years building distributed systems "
            "in Python and Go. Led migration of a monolithic Django app to a "
            "Kubernetes-based microservices architecture serving 50M MAUs."
        ),
        "skills": "Python, Go, FastAPI, Django, PostgreSQL, Redis, Kafka, Kubernetes, AWS, Terraform",
        "experience": (
            "Stripe — Staff Engineer (2021-Present): Led billing infra team of 6. "
            "Reduced 99p latency from 800ms to 120ms via async batching.\n"
            "Airbnb — Senior Engineer (2017-2021): Owned host payouts pipeline "
            "processing $4B annually."
        ),
        "education": "B.S. Computer Science, Carnegie Mellon University (2015)",
    },
    {
        "filename": "bob_chen.txt",
        "format": "txt",
        "name": "Bob Chen",
        "title": "Machine Learning Engineer",
        "email": "bob.chen@example.com",
        "phone": "+1-650-555-0193",
        "location": "Mountain View, CA",
        "summary": (
            "ML engineer specializing in NLP and large language model fine-tuning. "
            "Shipped a production RAG pipeline at Google serving 12M queries/day."
        ),
        "skills": "Python, PyTorch, TensorFlow, Hugging Face, LangChain, FAISS, BigQuery, GCP, Docker",
        "experience": (
            "Google — Senior MLE (2020-Present): Built domain-specific embedding "
            "models reducing search miss-rate by 23%.\n"
            "Cruise — MLE (2018-2020): Perception model training pipelines."
        ),
        "education": "M.S. Computer Science, Stanford University (2018)",
    },
    {
        "filename": "carol_davis.txt",
        "format": "txt",
        "name": "Carol Davis",
        "title": "Frontend Engineer",
        "email": "carol.davis@example.com",
        "phone": "+1-212-555-0157",
        "location": "New York, NY",
        "summary": (
            "Frontend engineer with deep React and TypeScript expertise. "
            "Built design systems used by 200+ engineers."
        ),
        "skills": "TypeScript, React, Next.js, Redux, Tailwind, Storybook, Cypress, Webpack, GraphQL",
        "experience": (
            "Spotify — Senior Frontend Engineer (2022-Present): Owns artist-tools "
            "web platform.\n"
            "Vercel — Frontend Engineer (2019-2022): Built dashboard analytics module."
        ),
        "education": "B.A. Computer Science, NYU (2019)",
    },
    {
        "filename": "david_kumar.txt",
        "format": "txt",
        "name": "David Kumar",
        "title": "DevOps / Platform Engineer",
        "email": "david.kumar@example.com",
        "phone": "+1-512-555-0181",
        "location": "Austin, TX",
        "summary": (
            "Platform engineer with 10 years in infra reliability. "
            "Reduced infra spend by 38% at scale via spot-instance and rightsizing."
        ),
        "skills": "Kubernetes, Terraform, AWS, GCP, Argo, Prometheus, Grafana, Python, Bash, Helm",
        "experience": (
            "Netflix — Senior SRE (2020-Present): Owns observability platform.\n"
            "HashiCorp — Engineer (2016-2020): Consul Connect mesh plane."
        ),
        "education": "B.E. Information Tech, IIT Bombay (2014)",
    },
    {
        "filename": "elena_petrov.docx",
        "format": "docx",
        "name": "Elena Petrov",
        "title": "Data Engineer",
        "email": "elena.petrov@example.com",
        "phone": "+1-206-555-0118",
        "location": "Seattle, WA",
        "summary": (
            "Data engineer focused on streaming ingestion and lakehouse "
            "architecture. Built Spark/Flink pipelines moving 12 TB/day."
        ),
        "skills": "Python, Scala, Spark, Flink, Kafka, Airflow, Snowflake, dbt, AWS, Databricks",
        "experience": (
            "Snowflake — Senior Data Engineer (2021-Present)\n"
            "Microsoft — Data Engineer (2017-2021): Azure Synapse ingestion team."
        ),
        "education": "M.S. Data Science, University of Washington (2017)",
    },
    {
        "filename": "frank_obrien.docx",
        "format": "docx",
        "name": "Frank O'Brien",
        "title": "Security Engineer",
        "email": "frank.obrien@example.com",
        "phone": "+1-617-555-0167",
        "location": "Boston, MA",
        "summary": (
            "Application security engineer with focus on AppSec, threat modeling, "
            "and SAST/DAST automation. OSCP and CISSP certified."
        ),
        "skills": "Python, Go, Burp, Semgrep, AWS, Kubernetes, OWASP, Threat Modeling, Pentesting",
        "experience": (
            "Cloudflare — Security Engineer (2020-Present): Owns customer-facing "
            "WAF rules.\n"
            "Akamai — AppSec Engineer (2017-2020)."
        ),
        "education": "B.S. Cybersecurity, Northeastern University (2017)",
    },
    {
        "filename": "grace_yamamoto.docx",
        "format": "docx",
        "name": "Grace Yamamoto",
        "title": "Product Manager (Technical)",
        "email": "grace.yamamoto@example.com",
        "phone": "+1-415-555-0124",
        "location": "San Francisco, CA",
        "summary": (
            "Technical PM with prior engineering background. Shipped 3 zero-to-one "
            "developer products generating $40M ARR."
        ),
        "skills": "Product Strategy, API Design, SQL, Python, A/B Testing, Roadmapping, Figma",
        "experience": (
            "Twilio — Senior PM (2021-Present): SendGrid Inbound Parse API.\n"
            "Atlassian — PM (2018-2021): Bitbucket Pipelines."
        ),
        "education": "M.B.A., Wharton (2018); B.S. CS, Brown (2014)",
    },
    {
        "filename": "henry_anderson.pdf",
        "format": "pdf",
        "name": "Henry Anderson",
        "title": "Full-Stack Engineer",
        "email": "henry.anderson@example.com",
        "phone": "+1-303-555-0149",
        "location": "Denver, CO",
        "summary": (
            "Full-stack engineer comfortable across the Python/TypeScript stack. "
            "Strong product sense, ex-founder."
        ),
        "skills": "Python, TypeScript, FastAPI, Next.js, PostgreSQL, Redis, Stripe, Docker",
        "experience": (
            "Notion — Software Engineer (2022-Present): Owns billing surface.\n"
            "Founder, ClipQuest (2019-2022): Acquired by Loom."
        ),
        "education": "B.S. Computer Science, University of Colorado Boulder (2019)",
    },
    {
        "filename": "isabella_rossi.pdf",
        "format": "pdf",
        "name": "Isabella Rossi",
        "title": "Site Reliability Engineer",
        "email": "isabella.rossi@example.com",
        "phone": "+1-310-555-0173",
        "location": "Los Angeles, CA",
        "summary": (
            "SRE with deep on-call experience at scale. Reduced p99 latency by 60% "
            "on a 200-service mesh."
        ),
        "skills": "Kubernetes, Istio, eBPF, Prometheus, Go, Python, Terraform, AWS, Chaos Engineering",
        "experience": (
            "Datadog — Senior SRE (2021-Present): Owns ingestion pipeline.\n"
            "Twitter — SRE (2018-2021): Timeline serving."
        ),
        "education": "B.S. Computer Engineering, UCLA (2018)",
    },
    {
        "filename": "james_williams.pdf",
        "format": "pdf",
        "name": "James Williams",
        "title": "Junior Software Engineer",
        "email": "james.williams@example.com",
        "phone": "+1-718-555-0166",
        "location": "Brooklyn, NY",
        "summary": (
            "Recent graduate with strong fundamentals in Python and web development. "
            "Two internships in fintech."
        ),
        "skills": "Python, Flask, React, SQL, Git, Linux, Algorithms, Data Structures",
        "experience": (
            "Plaid — SWE Intern (Summer 2024): Reduced webhook retry storms by 75%.\n"
            "Robinhood — SWE Intern (Summer 2023): Trading flow A/B test framework."
        ),
        "education": "B.S. Computer Science, Columbia University (2025)",
    },
]


def format_plain_text(r: dict[str, str]) -> str:
    return (
        f"{r['name']}\n"
        f"{r['title']}\n"
        f"{r['email']} | {r['phone']} | {r['location']}\n"
        "\n"
        "SUMMARY\n"
        f"{r['summary']}\n"
        "\n"
        "SKILLS\n"
        f"{r['skills']}\n"
        "\n"
        "EXPERIENCE\n"
        f"{r['experience']}\n"
        "\n"
        "EDUCATION\n"
        f"{r['education']}\n"
    )


def write_txt(target: Path, r: dict[str, str]) -> None:
    target.write_text(format_plain_text(r), encoding="utf-8")


def write_docx(target: Path, r: dict[str, str]) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(r["name"], level=0)
    title.alignment = 1  # center
    doc.add_paragraph(r["title"]).runs[0].bold = True
    doc.add_paragraph(f"{r['email']} | {r['phone']} | {r['location']}")

    for section in ("Summary", "Skills", "Experience", "Education"):
        doc.add_heading(section, level=2)
        doc.add_paragraph(r[section.lower()])
    doc.save(str(target))


def write_pdf(target: Path, r: dict[str, str]) -> None:
    """Produce a minimal valid PDF using pypdf — no system deps."""
    from io import BytesIO

    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        StreamObject,
        TextStringObject,
    )

    # Normalize unicode to a latin-1-safe ASCII subset; PDF Type1 Helvetica
    # uses single-byte encoding and would otherwise reject em-dashes, etc.
    replacements = {
        "—": "-",
        "–": "-",
        "‘": "'",
        "’": "'",
        "“": '"',
        "”": '"',
        "…": "...",
        " ": " ",
    }
    text = format_plain_text(r)
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    lines = text.splitlines() or [""]

    # Build a content stream that prints each line at a different y position.
    commands = ["BT", "/F1 11 Tf", "1 0 0 1 50 760 Tm"]
    for i, line in enumerate(lines):
        escaped = (
            line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        )
        if i == 0:
            commands.append(f"({escaped}) Tj")
        else:
            commands.append("0 -14 Td")
            commands.append(f"({escaped}) Tj")
    commands.append("ET")
    stream_data = "\n".join(commands).encode("latin-1")

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.pages[0]

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    resources = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_ref}
            )
        }
    )
    page[NameObject("/Resources")] = resources

    stream = StreamObject()
    stream._data = stream_data
    stream[NameObject("/Length")] = NumberObject(len(stream_data))
    stream_ref = writer._add_object(stream)
    page[NameObject("/Contents")] = stream_ref
    page[NameObject("/MediaBox")] = ArrayObject(
        [NumberObject(0), NumberObject(0), FloatObject(612), FloatObject(792)]
    )

    buf = BytesIO()
    writer.write(buf)
    target.write_bytes(buf.getvalue())


def main() -> None:
    out_dir = PROJECT_ROOT / "data" / "resumes"
    out_dir.mkdir(parents=True, exist_ok=True)
    for r in RESUMES:
        target = out_dir / r["filename"]
        fmt = r["format"]
        if fmt == "txt":
            write_txt(target, r)
        elif fmt == "docx":
            write_docx(target, r)
        elif fmt == "pdf":
            write_pdf(target, r)
        else:
            raise ValueError(f"unknown format {fmt}")
        print(f"wrote {target.relative_to(PROJECT_ROOT)}")


if __name__ == "__main__":
    main()
